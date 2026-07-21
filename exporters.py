"""Export a recording's transcript to various formats.

Text formats (txt, md, srt, vtt, json, csv, html) are pure Python - just
different serializations of the lines we already have. docx and pdf use
python-docx and reportlab, both pip-installable and PyInstaller-friendly,
so nothing extra is needed on the user's machine.

Every exporter takes the session dict {id, title, lines, summary} and
returns (bytes, media_type, file_extension).
"""

import csv
import io
import json


def _ts(seconds: float, sep: str = ",") -> str:
    """Seconds -> SRT/VTT timestamp HH:MM:SS,mmm (sep ',' for SRT, '.' VTT)."""
    ms = int(round(seconds * 1000))
    h, ms = divmod(ms, 3_600_000)
    m, ms = divmod(ms, 60_000)
    s, ms = divmod(ms, 1000)
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def _clock(seconds: float) -> str:
    s = int(seconds)
    if s >= 3600:
        return f"{s // 3600}:{(s % 3600) // 60:02d}:{s % 60:02d}"
    return f"{s // 60:02d}:{s % 60:02d}"


def _line_end(line, lines, idx):
    """Best-effort end time for a line: last word's end isn't stored, so use
    the next line's start, or start + a few seconds for the final line."""
    words = line.get("words") or []
    nxt = lines[idx + 1]["t"] if idx + 1 < len(lines) else None
    if nxt is not None:
        return max(line["t"], nxt - 0.05)
    # final line: estimate ~0.4s per word, min 2s
    return line["t"] + max(2.0, 0.4 * len(words) if words else 3.0)


# ------------------------------------------------------------ text formats
def to_txt(session):
    lines = session["lines"]
    body = "\n".join(f"[{_clock(l['t'])}] {l['text']}" for l in lines)
    header = session["title"] + "\n" + "=" * len(session["title"]) + "\n\n"
    return (header + body).encode("utf-8"), "text/plain", "txt"


def to_md(session):
    lines = session["lines"]
    out = [f"# {session['title']}", ""]
    if session.get("summary"):
        out += [session["summary"], "", "---", "", "## Transcript", ""]
    for l in lines:
        out.append(f"**`{_clock(l['t'])}`** {l['text']}")
        out.append("")
    return "\n".join(out).encode("utf-8"), "text/markdown", "md"


def to_srt(session):
    lines = session["lines"]
    out = []
    for i, l in enumerate(lines):
        out.append(str(i + 1))
        out.append(f"{_ts(l['t'])} --> {_ts(_line_end(l, lines, i))}")
        out.append(l["text"])
        out.append("")
    return "\n".join(out).encode("utf-8"), "application/x-subrip", "srt"


def to_vtt(session):
    lines = session["lines"]
    out = ["WEBVTT", ""]
    for i, l in enumerate(lines):
        out.append(f"{_ts(l['t'], '.')} --> {_ts(_line_end(l, lines, i), '.')}")
        out.append(l["text"])
        out.append("")
    return "\n".join(out).encode("utf-8"), "text/vtt", "vtt"


def to_json(session):
    payload = {
        "title": session["title"],
        "id": session["id"],
        "summary": session.get("summary"),
        "segments": [
            {"start": l["t"], "text": l["text"], "words": l.get("words", [])}
            for l in session["lines"]
        ],
    }
    return (json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8"),
            "application/json", "json")


def to_csv(session):
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["start_seconds", "timestamp", "text"])
    for l in session["lines"]:
        w.writerow([f"{l['t']:.2f}", _clock(l["t"]), l["text"]])
    return buf.getvalue().encode("utf-8"), "text/csv", "csv"


def _esc(s):
    return (s.replace("&", "&amp;").replace("<", "&lt;")
             .replace(">", "&gt;").replace('"', "&quot;"))


def to_html(session):
    lines = session["lines"]
    rows = "\n".join(
        f'<div class="line"><span class="ts">{_clock(l["t"])}</span>'
        f'<span class="tx">{_esc(l["text"])}</span></div>'
        for l in lines
    )
    summary_html = ""
    if session.get("summary"):
        summary_html = (f'<section class="summary"><h2>Summary</h2>'
                        f'<pre>{_esc(session["summary"])}</pre></section>')
    html = f"""<!DOCTYPE html>
<html lang="en"><head><meta charset="UTF-8">
<title>{_esc(session['title'])}</title>
<style>
  body{{font-family:system-ui,-apple-system,Segoe UI,sans-serif;
    max-width:760px;margin:40px auto;padding:0 20px;color:#1f2731;line-height:1.6}}
  h1{{font-size:24px}} h2{{font-size:16px;color:#8a6c2f}}
  .summary pre{{white-space:pre-wrap;background:#faf6ee;border:1px solid #e8dcc0;
    border-radius:8px;padding:14px;font-family:inherit}}
  .line{{margin-bottom:10px;display:flex;gap:10px}}
  .ts{{font-family:ui-monospace,Consolas,monospace;font-size:12px;color:#b8860b;
    min-width:52px;padding-top:2px}}
</style></head><body>
<h1>{_esc(session['title'])}</h1>
{summary_html}
<section><h2>Transcript</h2>{rows}</section>
</body></html>"""
    return html.encode("utf-8"), "text/html", "html"


# ------------------------------------------------------------ docx
def to_docx(session):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(session["title"], level=0)

    if session.get("summary"):
        doc.add_heading("Summary", level=1)
        for para in session["summary"].split("\n"):
            doc.add_paragraph(para)
        doc.add_heading("Transcript", level=1)

    for l in session["lines"]:
        p = doc.add_paragraph()
        run = p.add_run(f"[{_clock(l['t'])}]  ")
        run.bold = True
        run.font.color.rgb = RGBColor(0xB8, 0x86, 0x0B)
        run.font.size = Pt(9)
        p.add_run(l["text"])

    buf = io.BytesIO()
    doc.save(buf)
    return (buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx")


# ------------------------------------------------------------ pdf
def to_pdf(session):
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    body_style = ParagraphStyle("body", parent=styles["Normal"],
                                fontSize=10, leading=15, alignment=TA_LEFT)
    ts_style = ParagraphStyle("ts", parent=body_style,
                              textColor="#B8860B", fontName="Courier")

    story = [Paragraph(_esc(session["title"]), title_style), Spacer(1, 12)]

    if session.get("summary"):
        story.append(Paragraph("Summary", styles["Heading2"]))
        for para in session["summary"].split("\n"):
            if para.strip():
                story.append(Paragraph(_esc(para), body_style))
        story.append(Spacer(1, 10))
        story.append(Paragraph("Transcript", styles["Heading2"]))
        story.append(Spacer(1, 6))

    for l in session["lines"]:
        story.append(Paragraph(
            f'<font face="Courier" color="#B8860B">[{_clock(l["t"])}]</font> '
            + _esc(l["text"]), body_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue(), "application/pdf", "pdf"


EXPORTERS = {
    "txt": to_txt, "md": to_md, "srt": to_srt, "vtt": to_vtt,
    "json": to_json, "csv": to_csv, "html": to_html,
    "docx": to_docx, "pdf": to_pdf,
}
# Which formats are plain text (offer copy-to-clipboard, not just download).
TEXT_FORMATS = {"txt", "md", "srt", "vtt", "json", "csv", "html"}


def export(session, fmt):
    fn = EXPORTERS.get(fmt)
    if fn is None:
        raise ValueError(f"Unknown export format: {fmt}")
    return fn(session)